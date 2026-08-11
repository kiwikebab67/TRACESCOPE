import os
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Please run: pip install python-docx")
    exit(1)

doc = Document()

# Title
title = doc.add_heading('TraceScope DFIR Integration - Daily Progress Report', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 1. Deployment URL
doc.add_heading('1. Deployment URL', level=1)
p = doc.add_paragraph('Live Render Link: ')
p.add_run('https://tracescope-hq.onrender.com').bold = True

# 2. Measurable Outcomes
doc.add_heading('2. Measurable Outcomes (Today)', level=1)
doc.add_paragraph('• Modules Completed: 2 (Micro-Emulation Sandbox, Native Memory Forensics Engine)', style='List Bullet')
doc.add_paragraph('• Features Implemented: 4 (Speakeasy Detonation UI, ELA Heatmap Generator, Volatility 3 plugin execution, Local Ollama AI Payload De-obfuscator)', style='List Bullet')
doc.add_paragraph('• Files Modified/Created: 14+ files (Backend services for AI, emulation, and memory. Frontend UI panels in React).', style='List Bullet')
doc.add_paragraph('• Bugs Fixed: Resolved UI layout breaking on light-theme by migrating hardcoded Tailwind classes to global CSS theme variables.', style='List Bullet')

# 3. New Development vs. Refinements
doc.add_heading('3. New Development vs. Refinements', level=1)
p1 = doc.add_paragraph()
p1.add_run('New Development Completed Today:').bold = True
doc.add_paragraph('• Built and integrated the Volatility3 native Python engine to parse .raw and .mem files programmatically.', style='List Bullet')
doc.add_paragraph('• Integrated Speakeasy emulation pipeline to detonate PE files and hook API calls.', style='List Bullet')
doc.add_paragraph('• Created an offline AI De-obfuscation pipeline mapping to localhost:11434 (Llama3).', style='List Bullet')

p2 = doc.add_paragraph()
p2.add_run('Refinements / Optimizations from Earlier Work:').bold = True
doc.add_paragraph('• Optimized the frontend Evidence.jsx panel to enforce the TraceScope glassmorphic styling system.', style='List Bullet')

# 4. Challenges
doc.add_heading('4. Challenges Faced', level=1)
doc.add_paragraph('• Challenge: The Evidence UI had a button click failure issue due to overlapping absolute positioning, and the light theme was broken by hardcoded dark-mode elements.', style='List Bullet')
doc.add_paragraph('• Resolution: Re-architected the tab UI layout and CSS variables to seamlessly adapt to both themes and fixed z-index stacking issues.', style='List Bullet')

# 5. Pending Tasks
doc.add_heading('5. Pending Tasks & Planned Activities (Tomorrow)', level=1)
doc.add_paragraph('• Phase 3: Infrastructure & Telemetry integration.', style='List Bullet')
doc.add_paragraph('• Implement AWS CloudTrail parsing for Cloud forensics.', style='List Bullet')
doc.add_paragraph('• Build the eBPF Telemetry Hook simulated endpoints.', style='List Bullet')

# 6. Blockers
doc.add_heading('6. Blockers & Support Required', level=1)
doc.add_paragraph('No Blockers.', style='List Paragraph')

# Save file
downloads_dir = r"C:\Users\KARTIKEY\Downloads"
filepath = os.path.join(downloads_dir, 'TraceScope_Progress_Report.docx')
doc.save(filepath)
print(f"Report successfully generated at: {filepath}")
